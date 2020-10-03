import docx
d = docx.Document('c:\\Users\\Manith\\demo.docx')

p = d.paragraphs[0].text
p.runs[1].text #this should out put the bold style
p.runs[1].italic #changes to italic


p.style = 'Title' #change the style to title

d.save('c:\\Users\\Manith\\demo.docx')


#Make new doc
d2= docx.Document()
d2.add_paragraph('This is a new paragrpah')
p = d2.paragraphs[0]
p.add_run('New run')
d2.save('c:\\Users\\Manith\\demo2.docx')

def gettext(filename):
    doc = docx.Document(filename)
    fulltext = []
    for p in doc.paragraphs:
        fulltext.append(p.text)
    return '\n'.join(fulltext) #returns as string on new line